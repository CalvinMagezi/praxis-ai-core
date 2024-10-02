# tools/file_operations.py

import ell
from pathlib import Path
from ..workspace_manager import WorkspaceManager
from ..utils.logging import logger
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
import pypdf
import docx
import shutil
import mimetypes
import magic  # New dependency for improved file type detection

workspace_manager = WorkspaceManager()

def get_mime_type(file_path):
    """Get the MIME type of a file."""
    return magic.from_file(file_path, mime=True)

@ell.tool()
def read_file_tool(file_path: str, current_workspace: str) -> str:
    """
    Read the contents of a file in the specified workspace.

    Args:
    file_path (str): The path to the file within the workspace.
    current_workspace (str): The name of the current workspace.

    Returns:
    str: The contents of the file, or an error message if the file cannot be read.
    """
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    full_path = Path(workspace_path) / file_path
    try:
        mime_type = get_mime_type(str(full_path))
        if mime_type.startswith('text/'):
            with open(full_path, 'r', encoding='utf-8') as file:
                content = file.read()
        elif mime_type == 'application/pdf':
            content = read_pdf_tool(file_path, current_workspace)
        elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            content = read_word_document_tool(file_path, current_workspace)
        else:
            return f"Error: Unsupported file type: {mime_type}"
        return content
    except IOError as e:
        error_message = f"Error reading file: {full_path}. Error: {e}"
        logger.error(error_message)
        return error_message

@ell.tool()
def write_file_tool(file_path: str, content: str, current_workspace: str) -> str:
    """
    Write content to a file in the specified workspace.

    Args:
    file_path (str): The path to the file within the workspace.
    content (str): The content to write to the file.
    current_workspace (str): The name of the current workspace.

    Returns:
    str: A success message, or an error message if the file cannot be written.
    """
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    full_path = Path(workspace_path) / file_path
    try:
        full_path.parent.mkdir(parents=True, exist_ok=True)
        with open(full_path, 'w', encoding='utf-8') as file:
            file.write(content)
        success_message = f"File written successfully: {full_path}"
        logger.info(success_message)
        return success_message
    except IOError as e:
        error_message = f"Error writing file: {full_path}. Error: {e}"
        logger.error(error_message)
        return error_message

@ell.tool()
def create_folder_structure_tool(project_name: str, folder_structure: dict, code_blocks: list, current_workspace: str) -> str:
    """
    Create a folder structure with files in the specified workspace.

    Args:
    project_name (str): The name of the project (root folder).
    folder_structure (dict): A dictionary representing the folder structure.
    code_blocks (list): A list of tuples containing file names and their content.
    current_workspace (str): The name of the current workspace.

    Returns:
    str: A success message, or an error message if the structure cannot be created.
    """
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    project_path = Path(workspace_path) / project_name

    try:
        project_path.mkdir(parents=True, exist_ok=True)
        logger.info(f"Created project folder: {project_path}")
    except OSError as e:
        error_message = f"Error creating project folder: {project_path}. Error: {e}"
        logger.error(error_message)
        return error_message

    def create_folders_and_files(current_path: Path, structure: dict, code_blocks: list):
        for key, value in structure.items():
            path = current_path / key
            if isinstance(value, dict):
                try:
                    path.mkdir(exist_ok=True)
                    logger.info(f"Created folder: {path}")
                    create_folders_and_files(path, value, code_blocks)
                except OSError as e:
                    logger.error(f"Error creating folder: {path}. Error: {e}")
            else:
                code_content = next((code for file, code in code_blocks if file == key), None)
                if code_content:
                    try:
                        with open(path, 'w', encoding='utf-8') as file:
                            file.write(code_content)
                        logger.info(f"Created file: {path}")
                    except IOError as e:
                        logger.error(f"Error creating file: {path}. Error: {e}")
                else:
                    logger.warning(f"Code content not found for file: {key}")

    create_folders_and_files(project_path, folder_structure, code_blocks)
    return f"Folder structure created for project '{project_name}' in workspace '{current_workspace}'"

@ell.tool()
def create_pdf_tool(file_path: str, content: str, current_workspace: str) -> str:
    """Create a PDF file in the specified workspace."""
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    full_path = Path(workspace_path) / file_path
    try:
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        text_object = c.beginText(40, 750)
        for line in content.split('\n'):
            text_object.textLine(line)
        c.drawText(text_object)
        c.save()

        full_path.parent.mkdir(parents=True, exist_ok=True)
        with open(full_path, 'wb') as file:
            file.write(buffer.getvalue())
        return f"PDF file created successfully: {full_path}"
    except Exception as e:
        error_message = f"Error creating PDF file: {full_path}. Error: {e}"
        logger.error(error_message)
        return error_message

@ell.tool()
def read_pdf_tool(file_path: str, current_workspace: str) -> str:
    """Read the contents of a PDF file in the specified workspace."""
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    full_path = Path(workspace_path) / file_path
    try:
        with open(full_path, 'rb') as file:
            reader = pypdf.PdfReader(file)
            content = ""
            for page in reader.pages:
                content += page.extract_text()
        return content
    except Exception as e:
        error_message = f"Error reading PDF file: {full_path}. Error: {e}"
        logger.error(error_message)
        return error_message

@ell.tool()
def create_word_document_tool(file_path: str, content: str, current_workspace: str) -> str:
    """Create a Word document in the specified workspace."""
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    full_path = Path(workspace_path) / file_path
    try:
        doc = docx.Document()
        doc.add_paragraph(content)
        full_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(full_path)
        return f"Word document created successfully: {full_path}"
    except Exception as e:
        error_message = f"Error creating Word document: {full_path}. Error: {e}"
        logger.error(error_message)
        return error_message

@ell.tool()
def read_word_document_tool(file_path: str, current_workspace: str) -> str:
    """Read the contents of a Word document in the specified workspace."""
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    full_path = Path(workspace_path) / file_path
    try:
        doc = docx.Document(full_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return content
    except Exception as e:
        error_message = f"Error reading Word document: {full_path}. Error: {e}"
        logger.error(error_message)
        return error_message

@ell.tool()
def create_markdown_file_tool(file_path: str, content: str, current_workspace: str) -> str:
    """Create a Markdown file in the specified workspace."""
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    full_path = Path(workspace_path) / file_path
    try:
        full_path.parent.mkdir(parents=True, exist_ok=True)
        with open(full_path, 'w', encoding='utf-8') as file:
            file.write(content)
        return f"Markdown file created successfully: {full_path}"
    except Exception as e:
        error_message = f"Error creating Markdown file: {full_path}. Error: {e}"
        logger.error(error_message)
        return error_message

@ell.tool()
def read_markdown_file_tool(file_path: str, current_workspace: str) -> str:
    """Read the contents of a Markdown file in the specified workspace."""
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    full_path = Path(workspace_path) / file_path
    try:
        with open(full_path, 'r', encoding='utf-8') as file:
            content = file.read()
        return content
    except Exception as e:
        error_message = f"Error reading Markdown file: {full_path}. Error: {e}"
        logger.error(error_message)
        return error_message

@ell.tool()
def copy_file_tool(source_path: str, destination_path: str, current_workspace: str) -> str:
    """Copy a file within the specified workspace."""
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    source_full_path = Path(workspace_path) / source_path
    destination_full_path = Path(workspace_path) / destination_path

    try:
        destination_full_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source_full_path, destination_full_path)
        return f"File copied successfully from {source_full_path} to {destination_full_path}"
    except Exception as e:
        error_message = f"Error copying file from {source_full_path} to {destination_full_path}. Error: {e}"
        logger.error(error_message)
        return error_message

@ell.tool()
def move_file_tool(source_path: str, destination_path: str, current_workspace: str) -> str:
    """Move a file within the specified workspace."""
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    source_full_path = Path(workspace_path) / source_path
    destination_full_path = Path(workspace_path) / destination_path

    try:
        destination_full_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(source_full_path, destination_full_path)
        return f"File moved successfully from {source_full_path} to {destination_full_path}"
    except Exception as e:
        error_message = f"Error moving file from {source_full_path} to {destination_full_path}. Error: {e}"
        logger.error(error_message)
        return error_message

@ell.tool()
def delete_file_tool(file_path: str, current_workspace: str) -> str:
    """Delete a file from the specified workspace."""
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    full_path = Path(workspace_path) / file_path

    try:
        full_path.unlink()
        return f"File deleted successfully: {full_path}"
    except Exception as e:
        error_message = f"Error deleting file: {full_path}. Error: {e}"
        logger.error(error_message)
        return error_message

@ell.tool()
def list_files_tool(directory_path: str, current_workspace: str) -> str:
    """List files in a directory within the specified workspace."""
    workspace_path = workspace_manager.get_workspace_path(current_workspace)
    if not workspace_path:
        return f"Error: No valid workspace path for workspace: {current_workspace}"

    full_path = Path(workspace_path) / directory_path

    try:
        files = [f.name for f in full_path.iterdir() if f.is_file()]
        return f"Files in {full_path}:\n" + "\n".join(files)
    except Exception as e:
        error_message = f"Error listing files in directory: {full_path}. Error: {e}"
        logger.error(error_message)
        return error_message